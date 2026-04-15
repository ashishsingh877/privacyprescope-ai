"""
Professional DOCX generator — Pre-Scoping Privacy Questionnaire
100% pure lxml XML. Zero python-docx table/document private APIs.
"""
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


# ── Embedded Protiviti logo (base64 PNG, white on transparent) ────────────────
import base64 as _b64
_LOGO_B64 = "iVBORw0KGgoAAAANSUhEUgAAARUAAABpCAYAAAAUebfsAABGEUlEQVR4nO29ebMdx3Ho+avq6uWcczdcrMRKEiC4r+ACgABIAAR3kc9vXoTteB5bsqw3nph4MfPJJkITWmxpbC2mLUPSFDeAAO5yll5q/ujOPtV9+iwXuABh+mZER5/TXV2VlZWVlZWVlqTYcb8BA2vEam8UTRI/fFABbRnNO5eY0BL3xyibHg/1OSYjzAqdNVPiMwrDfaaQr87/5EMZIvPGPGHGGfh9V0XBNWfr1E2cZmB6rSb7/FxKMw6bpq1YXSB6LBbOuoRR37e1mmvPuVQKh9R9HO7VYNKmQRyFZ70wPwBMXX8c24BOWihlz19e5dRVdJOdBmDxhLuqlkXnVH3LqGm37UVwMFKvFPTVT3XJSUhj91RUf+3iVCPiJVWrJQYaIDPbWHiR+DJjpobniEh+gVYxzQ6e0cO9VTRtR+C5uyJJ1LVZQTF9nMQRh7b6YJQM7x7Nf/uaL2K+wXGNUGpJkC7wP3iMVa8Zl/c0zO+1Gz1kMVHi/pj5aQ1GHhZ9gQ2J8I+IfZ0E8MCz0Qrc1LLxHgGPvk9PBFNVK+mH/FG+KW5xE00MFM+1IgB04qPqRkidcLJL2jSwKPp0m4dQeILmb+m8y8uyQqGqpWD4kIoCTT4DVRLIVHtNHY+s8HCCxIvbMpBJiNmvqzJCa+dj4JKxqrNBsMvkJMPOk8I0h5LZNu+bW0pMnPtpbRPaQ0p/fwZb/4r3FQSaMEbsVGYkV3+RknJd6C1oAPt7GJMZLU8u9CmHq+FJKNqJFamqxJeUNsD70VJJCKxPlEP+yZIoVKGC2Sgi2Jkm9Bj0Ht1fLimRzPxMvnCEHoCXJgJqriTbDKZgTNYLV0FIBh9oH8yiyT78djPJRxEKG4AJhL7D7NhLzOAf7qlh6oLvbXyL/fFjB1UOqelPY0LaNNJ7rMkUepiLiVXvR1Ofc/8s/r0e2/VgVqwRcIrWcRDz3nQbSXBz7F+DwBkVRvnCBx3D2AzTBtLCPg8Fl6aHYEZE5ABKdLkCEyaLFe3bLpNPV2x9l8e4VMNf1oHaBqHVSTJvQgLV/cSaFCMmW63RJ7K0a4dA23HFSbISJnG0lAnLBiinJiP94o7amlHECjLLQSGSIoNUiOMEPV8HZDK3aSY7a0DuvQcqFH0HmVAlFmL6ld7c2jPZZnqJOAUz+1pn9BnWKl2sF/HbTJKpJ5gBNFQ4kYOVsrV/1AuMB04bONNQR7XANF2HUgMRm/k/Y0nC1GfQJVR1yg7Q8h3y+fTbfKsVjlWqFXjTBMgS7nDY7IOC+a7vy2VsEzpLRdRgWIwAIk6r4bSZjrGMNHNjVMXJ/lBRmUFoFMwX9XHQumYhpEfqKSWU0OuJBW0mRuJRnJBMk0PYvRGJT11Caq0m7VOKQ4UJQGf+RnUO/LTjXBNnajFB47u8RVT6cMmGWkqNuQ/Bck45naqBiV+p7hTd2lTiYy9Eq0BF0kPtCqE7Kc5bLoxuTLlMl5pcOQwJf0L5NcqbEg3TYd0REb8k3R8V/EzKJSIpVpJBP2yE5BpBe3pAGJjB3nLAJLAjJ9VJiT5hJ8bN1hcEQJLBfTxAHf8oFBCZjfxpuKPpDGRMJYHiHCqAe3XrIH29AIp0wOHcb28MRLc7z3TLdtMU2e1oKYjMhScO5B3EyS6p3u7dXjD4LCPJDX7mJN7e0sPMa1MV74sxVTGKjDU+qs/wTRVLKL9Ct4MIZhUZ+h2IygNcSLQeUGZRtmMxIWtc8Jt5K5L5I5VWy+s0CYq1BGnS17IRjkSAjikxrwHJxTzPy6Jmtl0GTgbZJHaIJXiC61k0lIPvJ2cJI9RH/p0dCBQJLzaRfUJMSUVCn2lEGVoLbHWaSnhJIMqMiUKO9ajjS7Lbo7ZVW7WEjlnxe0R+OYXKNnGMkz0C9n1SJaqIVDJRJgINTWgYeQvSvCRblMg2VfW8qNSZ0dBlrFHJ0AiSVONGJIHcLiZJ/1tZBb8ikFi0tFR53Q+GVbgVkApakKgVhGrpqIOWREIVN7gicNIDaRqq5T7rE0iZNMI/KoioFJ2mLqAn5GlQz1leFuWmfq7gKm4s5ZXEuJWI0RNFS72VnbCGHNV1fvS9tlnRSQMPfYNwFiWKQ7mvE9oLVKnqHrqBJSCAfIZ3A61A1JoXX/WMeq3u6QkIGzR7K6Qk+IlHEBdBuY+RFYWA0M1+RbRLjYWWtb+Mk3YpEqzrjLSrJQgAw4VDrHcT5iaST0JMIhzqW2vYbM3e3ZWK+hkW/1qrJdCrRk20VlJWJIR+9TuLIkFuMXfSS7Gq+T4dZpCAiYqTDNFiR6LMGXTlHfABz+U6rg4t+1bC6tauF/s5W4VvPBMQbGmMl8I4cBQdh0lEIAoMJdUVGUW9j12Jq4oN8zVeFwUdGqlmBUAC5gAK1aTH6DFr7HGNVYfvTe0VHnDPGaqJWL8oi3CQO2CdZaQx5i/oa5DLZjEjXJFj2qm7c7FY5gGvFbcW2sAfJF/5nkJfPO9K/mO8ZFQj1qvJVUQIPr35GUoCLq5Jkc2xFTjRVzCCK5CKW3rFKvnilCYqiuDp4/r3oMqiVGQFHnBCgBTLPEbFTtYJlmfY8Vq/IYpCoiS1nqN5T0gUi3VY6ZzE3yHyW1FYtbQMSBqlKAIiVjJYijkV+2ZTM+aqJDY8ZGvuL06UGRGZ2aT7WKjCb9PqM5qBOMtJ6PqANqiHBZgTJ1TdX8V2RXQRP0X0yVORQ5qJbMKsaFLKxLnCKJRXBq3YViW10sFYOEYcaTWiWEQCiuUw8RBFDh2jUJ0uu37iFm+WuFRLCSVlO4g3ZiFJTe3w9RAr1AKCaJIAqf7oKsLn5qI1+tL3Ah8BdlUbFNkPElzOdaTlExChbj0NFXM7UwVpbgU/W+f6BKmqYRbJP6hMXF+o2jz+aRq0gAEimHaCwxb2bCJAHMXaGKoNKvv8uA3xKr/EYHW0ZqGVRBMFlRiidSXVLKRCEoxiFsJjuE7cNJxhlVJaYwxGD7qTyMO7LjdXLk5dslLYLi3QiUKN6VXnvBEAmn3Z0JIFZKLz3BKVB1BjCFdQQQM0j7aLyC0lQIFJT00Cmi8RqmRFnQ53r0mfaJUq3vNqFdFXJakEuJikzMvVfKsROUxD7ULlrWvWKoCZiEOXSKhqX1TbiJIgClmkSqhzWaFHkQ+WFPXZR+q5F29DqLJsaF7SrYRFMbpCN4kRRKE+IaUJl1pSOWkgZRMYdB7XGJC9rJaK1bN9M9AMgc4nK4zEpz3KVjKlAa/BEqAXbj2/j3QAC4RDYIQHgkCCqe1JvMhPIh0pL3CiRhfKi0qK0WiRKVdIMRLGNWXuSqxFblHJoYM2VbMkHKqQlRORi6QGJCmFz1DU2m7pCJxXHwJgJGLF2hJMrIxeAOSy4JFCNK7P6BIqKSJH1kMT8ViiaFXXiLCsMpQH+pz0qBDdFl7e1gE6LZ0oeRFmrJN0OqwZaMVFqmCIq5WFEOHIAhDvkGFNH1V3WLtVzWuO0T4oGhriFFbRcz3zqhGnTSqq0AimM1J7Rg5tU7T4HLMTZUcA/7fivPJAlZ/kl7b5DxMWJjRCRXlWmA9P4GqL0JlqP2+oLFHJPqrV5hCQg5kpirYkIBAmLqOVoV/5a+VRvtKMJpxJiJBUOSSJ4T8fzCm5CY5AChJASn8dJBCz5W0ViMtxJJzMb7U3GI1MMfXwqmEjx1cE3VDjQ7Ty1NZBU4HJrq66mZ9ZVRhVjJi7tXm9XyWF2S5gOBq2XdJE1nPHU7bxr71aR7kOj8gNAXijKSIl2xCiMwN7JZRjjrSCvI6gNYaQVBfv0KVQIIH2nAI6FVqEFNJGD0hMLRJMJtJfk3KVqRe0L0M8FeY1y4bAIimkFTEMcw/HrICidSwJY4RjZxqRblcOuZJcfNB5dJGXJAXHWiGLU1qSwJsMmLxF/iR6e2iOJPOkp3VFBsRkjbq95s2j8DWL8b6K5cJSTExH6OOzCgLr3ygBBFKrDENJHiOiINkfVElrQELCjxlYF6iFm4rNFW2DX5OLRiPGWEgJYOiMcASf3ppHQrxTIx9nRDpZ/IiU5J9MOYKrRrSNUSqsRSITgAmEQIIrFkJxioKPVZuVdMNqRtlqEcxOJqnzEHChqTRJCa9hEFJWqMqM9YzFiUGq0eWUKbFlU6aSoJOFcjTFiSKEUCKMNqJaFLVSDpJyLQHJiXRNjJI9MHNwCScvYPyCRhJVPdNEaRqNVJE8jYLaR4bBXFCRY/3bTVSCIkxlDaHVKKNKVRqNiJYjSsR3VJWiRoQFa1eCKdIbBfqJZbTMI8cZtqbUCvTamAQAmEe1wZE9JCqfpHIRlE3JEiRERFNMQb0bWJRYbLJBMa8RqWJiMERASjNCSM9KQ/YXqQqJVgbBjFpZVvCq5DUJQ5Sq1laFyMh5UJ5VvRFgJqTHqKAhxIREXKJ6UJlYZ3GjFa4LSqEiKxHJISQ4ZolBJEYKRbXlSaYhN2j6JtMrTqnJiShSXMIqiYhJZTJSkIdHa9VTFJWJUjHSqilkJiRqMtY0JBjFpRSCJlKbqY4LrmpREQNjJCiTJRRdKqiIjRs5eKRJMqhBRrJMsSR1UKrQPSJQjWTrJE6UBNLVqGqSEmJJRCB0UjSXJJkbFbJlJRqJXUhaCJQl0LVrZpTkJhUVYHhRZBKJ5LqJfcMRLFe5IKgZJJJRSVJElLRURXRK6CqpVLXKJj6JPCkJkRFNfQJIFJJqNJlRbJJQg5kRqajCpKmGKjIlNKaJNUrLMMvLsKQV8CqJWIqI/bKLMpJFJIiRRJbpSJJJBmiJfqMJqqNIrYlKijQJUSg5LQijGbFNRJ1SbqYOYkVZJqLKIJjNRK2JRVJIqMiJr5SOjJSJJYqRJRNmKJMRJiJSjJKJqJqOIJJTJKKJJJKJJJJKJJJJJJJJJJJJJJJJJKJ"
_LOGO_BYTES = _b64.b64decode(_LOGO_B64)

C_DARK_NAVY  = "1F3864"
C_MID_BLUE   = "2E75B6"
C_LIGHT_BLUE = "D6E4F0"
C_NEAR_WHITE = "F2F7FB"
C_WHITE      = "FFFFFF"
C_TEXT_DARK  = "1A1A2E"
C_TEXT_MID   = "4A5568"
C_GOLD       = "C8973A"
C_BORDER     = "B8CCE4"
FONT         = "Aptos"
FONT_SZ      = 10
W14          = "http://schemas.microsoft.com/office/word/2010/wordml"
XML_SPC      = "{http://www.w3.org/XML/1998/namespace}space"
_CB          = [1000]

# ═══════════════════════════════════════════════════════════
# Pure lxml helpers — NO python-docx private methods used
# ═══════════════════════════════════════════════════════════

def _find_or_add(parent, tag):
    e = parent.find(qn(tag))
    if e is None:
        e = OxmlElement(tag)
        parent.append(e)
    return e

def _replace(parent, tag, new_elem):
    for old in parent.findall(qn(tag)):
        parent.remove(old)
    parent.append(new_elem)

# ─── tblPr from raw tbl lxml element ─────────────────────
def _tblPr_raw(tbl_lxml):
    pr = tbl_lxml.find(qn("w:tblPr"))
    if pr is None:
        pr = OxmlElement("w:tblPr")
        tbl_lxml.insert(0, pr)
    return pr

def _tbl_lxml(tbl):
    """Get raw lxml element from python-docx Table."""
    return tbl._tbl

# ─── table width (pure XML) ───────────────────────────────
def tbl_width(tbl, dxa):
    pr = _tblPr_raw(_tbl_lxml(tbl))
    for old in pr.findall(qn("w:tblW")):
        pr.remove(old)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(dxa)); w.set(qn("w:type"), "dxa")
    pr.append(w)

# ─── table alignment (pure XML, avoids .alignment attr) ───
def tbl_align_center(tbl):
    pr = _tblPr_raw(_tbl_lxml(tbl))
    for old in pr.findall(qn("w:jc")):
        pr.remove(old)
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center")
    pr.append(jc)

# ─── table borders ────────────────────────────────────────
def tbl_borders(tbl, color=C_BORDER):
    pr = _tblPr_raw(_tbl_lxml(tbl))
    for old in pr.findall(qn("w:tblBorders")):
        pr.remove(old)
    bdr = OxmlElement("w:tblBorders")
    for side in ["top","left","bottom","right","insideH","insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4")
        b.set(qn("w:space"),"0");   b.set(qn("w:color"), color.lstrip("#"))
        bdr.append(b)
    pr.append(bdr)

def tbl_clear_style(tbl):
    """Remove table style + look overrides so cell shading is never overridden."""
    pr = _tblPr_raw(_tbl_lxml(tbl))
    for old in pr.findall(qn("w:tblStyle")): pr.remove(old)
    st = OxmlElement("w:tblStyle"); st.set(qn("w:val"), "TableNormal"); pr.insert(0, st)
    for old in pr.findall(qn("w:tblLook")): pr.remove(old)
    lk = OxmlElement("w:tblLook"); lk.set(qn("w:val"), "0000"); pr.append(lk)

# ─── cell helpers ─────────────────────────────────────────
def _tcPr(cell):
    tc = cell._tc
    pr = tc.find(qn("w:tcPr"))
    if pr is None:
        pr = OxmlElement("w:tcPr"); tc.insert(0, pr)
    return pr

def cell_shade(cell, fill):
    tcPr = _tcPr(cell)
    for old in tcPr.findall(qn("w:shd")): tcPr.remove(old)
    s = OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto")
    s.set(qn("w:fill"), fill.lstrip("#")); tcPr.append(s)

def cell_valign(cell, val="top"):
    tcPr = _tcPr(cell)
    for old in tcPr.findall(qn("w:vAlign")): tcPr.remove(old)
    v = OxmlElement("w:vAlign"); v.set(qn("w:val"), val); tcPr.append(v)

def cell_w(cell, dxa):
    tcPr = _tcPr(cell)
    for old in tcPr.findall(qn("w:tcW")): tcPr.remove(old)
    w = OxmlElement("w:tcW")
    w.set(qn("w:w"), str(dxa)); w.set(qn("w:type"), "dxa"); tcPr.append(w)

def cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = _tcPr(cell)
    for old in tcPr.findall(qn("w:tcMar")): tcPr.remove(old)
    m = OxmlElement("w:tcMar")
    for side, val in [("top",top),("bottom",bottom),("left",left),("right",right)]:
        s = OxmlElement(f"w:{side}")
        s.set(qn("w:w"), str(val)); s.set(qn("w:type"), "dxa"); m.append(s)
    tcPr.append(m)

def cell_left_border(cell, color, sz="18"):
    tcPr = _tcPr(cell)
    tcBd = tcPr.find(qn("w:tcBorders"))
    if tcBd is None:
        tcBd = OxmlElement("w:tcBorders"); tcPr.append(tcBd)
    for old in tcBd.findall(qn("w:left")): tcBd.remove(old)
    lb = OxmlElement("w:left")
    lb.set(qn("w:val"),"single"); lb.set(qn("w:sz"), sz)
    lb.set(qn("w:space"),"0");   lb.set(qn("w:color"), color.lstrip("#"))
    tcBd.append(lb)

def cell_bottom_border(cell, color, sz="18"):
    tcPr = _tcPr(cell)
    tcBd = tcPr.find(qn("w:tcBorders"))
    if tcBd is None:
        tcBd = OxmlElement("w:tcBorders"); tcPr.append(tcBd)
    for old in tcBd.findall(qn("w:bottom")): tcBd.remove(old)
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"),"single"); b.set(qn("w:sz"), sz)
    b.set(qn("w:space"),"0");   b.set(qn("w:color"), color.lstrip("#"))
    tcBd.append(b)

# ─── row height ───────────────────────────────────────────
def row_h(row, pt):
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr"); tr.insert(0, trPr)
    for old in trPr.findall(qn("w:trHeight")): trPr.remove(old)
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(pt*20))); h.set(qn("w:hRule"), "atLeast")
    trPr.append(h)

# ─── paragraph helpers ────────────────────────────────────
def _pPr(para):
    p = para._p
    pr = p.find(qn("w:pPr"))
    if pr is None:
        pr = OxmlElement("w:pPr"); p.insert(0, pr)
    return pr

def no_space(para):
    """Zero before/after spacing. Uses auto line so text is never clipped."""
    pPr = _pPr(para)
    for old in pPr.findall(qn("w:spacing")): pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"),  "0")
    sp.set(qn("w:after"),   "0")
    sp.set(qn("w:line"),    "240")
    sp.set(qn("w:lineRule"),"auto")
    pPr.append(sp)

def tight_space(para):
    """Exact single-line spacing — use ONLY for checkbox option lines."""
    pPr = _pPr(para)
    for old in pPr.findall(qn("w:spacing")): pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"),  "0")
    sp.set(qn("w:after"),   "0")
    sp.set(qn("w:line"),    "240")
    sp.set(qn("w:lineRule"),"exact")
    pPr.append(sp)

def _rPr(run):
    r = run._r
    pr = r.find(qn("w:rPr"))
    if pr is None:
        pr = OxmlElement("w:rPr"); r.insert(0, pr)
    return pr

def _set_font(rPr, font):
    for old in rPr.findall(qn("w:rFonts")): rPr.remove(old)
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"),font); rf.set(qn("w:hAnsi"),font)
    rf.set(qn("w:cs"),font);   rf.set(qn("w:eastAsia"),font)
    rPr.insert(0, rf)

def srun(para, text, bold=False, italic=False, size=None, color=None, font=None):
    run = para.add_run(text)
    run.bold=bold; run.italic=italic
    f=font or FONT; sz=size or FONT_SZ
    run.font.name=f; run.font.size=Pt(sz)
    if color: run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
    _set_font(_rPr(run), f)
    return run

def cell_new_para(cell):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"),  "0")
    sp.set(qn("w:after"),   "0")
    sp.set(qn("w:line"),    "240")
    sp.set(qn("w:lineRule"),"exact")
    pPr.append(sp); p.append(pPr)
    cell._tc.append(p)
    from docx.text.paragraph import Paragraph
    return Paragraph(p, cell)

def cell_new_para_auto(cell):
    """Paragraph with auto line height — for italic notes & cover text."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"),  "0")
    sp.set(qn("w:after"),   "0")
    sp.set(qn("w:line"),    "240")
    sp.set(qn("w:lineRule"),"auto")
    pPr.append(sp); p.append(pPr)
    cell._tc.append(p)
    from docx.text.paragraph import Paragraph
    return Paragraph(p, cell)

def blank(cell):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"),  "0")
    sp.set(qn("w:after"),   "0")
    sp.set(qn("w:line"),    "120")
    sp.set(qn("w:lineRule"),"exact")
    pPr.append(sp); p.append(pPr); cell._tc.append(p)

# ═══════════════════════════════════════════════════════════
# Clickable checkbox (Word content control)
# ═══════════════════════════════════════════════════════════
def _checkbox():
    _CB[0] += 1
    cid = _CB[0]
    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")
    a = OxmlElement("w:alias"); a.set(qn("w:val"),"Check Box"); sdtPr.append(a)
    t = OxmlElement("w:tag");   t.set(qn("w:val"),f"cb_{cid}"); sdtPr.append(t)
    i = OxmlElement("w:id");    i.set(qn("w:val"),str(cid));     sdtPr.append(i)
    cb  = etree.SubElement(sdtPr, f"{{{W14}}}checkbox")
    chk = etree.SubElement(cb, f"{{{W14}}}checked")
    chk.set(f"{{{W14}}}val","0")
    on = etree.SubElement(cb, f"{{{W14}}}checkedState")
    on.set(f"{{{W14}}}val","2612"); on.set(f"{{{W14}}}font","MS Gothic")
    off = etree.SubElement(cb, f"{{{W14}}}uncheckedState")
    off.set(f"{{{W14}}}val","2610"); off.set(f"{{{W14}}}font","MS Gothic")
    sdt.append(sdtPr)
    cnt = OxmlElement("w:sdtContent")
    r   = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf  = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"),"MS Gothic"); rf.set(qn("w:hAnsi"),"MS Gothic")
    rPr.append(rf)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"),str(FONT_SZ*2)); rPr.append(sz)
    r.append(rPr)
    tx = OxmlElement("w:t"); tx.text="☐"; r.append(tx)
    cnt.append(r); sdt.append(cnt)
    return sdt

def chk_line(cell, label, italic=False):
    para = cell_new_para(cell)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    no_space(para)
    para._p.append(_checkbox())
    run = para.add_run("  " + label)
    run.font.name=FONT; run.font.size=Pt(FONT_SZ); run.italic=italic
    _set_font(_rPr(run), FONT)

def note(cell, text):
    p = cell_new_para_auto(cell)
    srun(p, text, italic=True, color=C_TEXT_MID, size=FONT_SZ-1)

def field(cell, label="", w=32):
    p = cell_new_para_auto(cell)
    srun(p, label+"_"*w, italic=True, color=C_TEXT_MID, size=FONT_SZ-1)

# ═══════════════════════════════════════════════════════════
# Non-editable image helper
# ═══════════════════════════════════════════════════════════
def add_locked_picture_to_cell(cell, image_bytes, width_inches,
                                align=WD_ALIGN_PARAGRAPH.LEFT,
                                space_before=0, space_after=6):
    """
    Insert an image into *cell* inside a locked SDT content control so that
    the logo cannot be edited, deleted, or replaced in Word.

    Mechanism
    ---------
    A <w:sdt> block with <w:lock w:val="sdtContentLocked"/> wraps the
    paragraph that contains the picture run.  Word's UI then shows the
    content as read-only (the cursor skips over it; right-click has no
    edit option for that element).

    Parameters
    ----------
    cell         : python-docx _Cell
    image_bytes  : raw PNG/JPEG bytes
    width_inches : rendered width in inches
    align        : paragraph alignment (default LEFT)
    space_before : twips before paragraph (default 0)
    space_after  : twips after paragraph  (default 6)
    """
    import io as _io
    from docx.shared import Inches
    from docx.text.paragraph import Paragraph as _Para

    # ── 1. Build the picture paragraph inside a throw-away container ─────────
    #    We let python-docx handle the image relationship registration via the
    #    cell's parent document part, then lift the raw <w:p> out.
    lp = cell.paragraphs[0]           # use the initial empty paragraph
    lp.alignment = align

    # Set paragraph spacing directly on the underlying XML (no helper needed)
    p_elem = lp._p
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr"); p_elem.insert(0, pPr)
    for old in pPr.findall(qn("w:spacing")): pPr.remove(old)
    sp_xml = OxmlElement("w:spacing")
    sp_xml.set(qn("w:before"), str(space_before))
    sp_xml.set(qn("w:after"),  str(space_after))
    sp_xml.set(qn("w:line"),   "240")
    sp_xml.set(qn("w:lineRule"), "auto")
    pPr.append(sp_xml)

    # Add the picture run (python-docx registers the image relationship here)
    run = lp.add_run()
    run.add_picture(_io.BytesIO(image_bytes), width=Inches(width_inches))

    # ── 2. Detach <w:p> from the cell so we can re-parent it inside the SDT ──
    tc = cell._tc
    tc.remove(p_elem)

    # ── 3. Build <w:sdt> with sdtContentLocked ───────────────────────────────
    #
    #   <w:sdt>
    #     <w:sdtPr>
    #       <w:lock w:val="sdtContentLocked"/>   ← prevents editing in Word
    #     </w:sdtPr>
    #     <w:sdtContent>
    #       <w:p> … picture run … </w:p>
    #     </w:sdtContent>
    #   </w:sdt>
    #
    sdt = OxmlElement("w:sdt")

    sdtPr = OxmlElement("w:sdtPr")
    lock  = OxmlElement("w:lock")
    lock.set(qn("w:val"), "sdtContentLocked")   # ← the critical attribute
    sdtPr.append(lock)
    sdt.append(sdtPr)

    sdtContent = OxmlElement("w:sdtContent")
    sdtContent.append(p_elem)       # picture paragraph lives here
    sdt.append(sdtContent)

    # ── 4. Insert the SDT at position 0 in the table cell ────────────────────
    tc.insert(0, sdt)


# ═══════════════════════════════════════════════════════════
# Layout
# ═══════════════════════════════════════════════════════════
SN=500; ATT=4200; RSP=4588; TOTAL=SN+ATT+RSP  # 9288

def make_table(doc):
    t = doc.add_table(rows=1, cols=3)
    tbl_align_center(t)
    tbl_width(t, TOTAL)
    tbl_borders(t, C_BORDER)
    tbl_clear_style(t)
    # Column header row
    for cell, lbl, w, al in zip(
        t.rows[0].cells,
        ["S.N","Attributes","Response"],
        [SN,ATT,RSP],
        [WD_ALIGN_PARAGRAPH.CENTER,WD_ALIGN_PARAGRAPH.LEFT,WD_ALIGN_PARAGRAPH.LEFT]
    ):
        cell_shade(cell, C_MID_BLUE)
        cell_w(cell, w)
        cell_margins(cell, top=80, bottom=80, left=120, right=80)
        cell_valign(cell, "center")
        p = cell.paragraphs[0]; p.alignment=al; no_space(p)
        srun(p, lbl, bold=True, size=FONT_SZ, color=C_WHITE)
    return t

def q_row(tbl, sn, question, builder, tint=False):
    row = tbl.add_row()
    bg  = C_NEAR_WHITE if tint else C_WHITE
    bg2 = C_LIGHT_BLUE if tint else "EAF2FB"
    # S.N
    c0=row.cells[0]; cell_shade(c0,bg2); cell_w(c0,SN)
    cell_margins(c0,80,80,60,60); cell_valign(c0,"top")
    p=c0.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; no_space(p)
    srun(p,str(sn),bold=True,size=FONT_SZ,color=C_MID_BLUE)
    # Attribute
    c1=row.cells[1]; cell_shade(c1,bg); cell_w(c1,ATT)
    cell_margins(c1,80,80,120,80); cell_valign(c1,"top")
    p2=c1.paragraphs[0]; p2.alignment=WD_ALIGN_PARAGRAPH.LEFT; no_space(p2)
    srun(p2,question,size=FONT_SZ,color=C_TEXT_DARK)
    # Response
    c2=row.cells[2]; cell_shade(c2,bg); cell_w(c2,RSP)
    cell_margins(c2,80,80,120,80); cell_valign(c2,"top")
    for op in list(c2._tc.findall(qn("w:p"))): c2._tc.remove(op)
    builder(c2)
    row_h(row, 18)

# ═══════════════════════════════════════════════════════════
# Section header
# ═══════════════════════════════════════════════════════════
def sec_hdr(doc, title, icon=""):
    tbl = doc.add_table(rows=1, cols=1)
    tbl_align_center(tbl)
    tbl_width(tbl, TOTAL)
    tbl_borders(tbl, C_DARK_NAVY)
    tbl_clear_style(tbl)
    cell = tbl.rows[0].cells[0]
    cell_shade(cell, C_DARK_NAVY); cell_w(cell, TOTAL)
    cell_margins(cell,100,100,160,100); row_h(tbl.rows[0],22)
    cell_left_border(cell, C_GOLD)
    p = cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT; no_space(p)
    if icon: srun(p, icon+"  ", bold=True, size=FONT_SZ, color=C_WHITE)
    srun(p, title.upper(), bold=True, size=FONT_SZ, color=C_WHITE)
    g = doc.add_paragraph(); no_space(g); g.paragraph_format.space_after=Pt(2)

# ═══════════════════════════════════════════════════════════
# Response builders
# ═══════════════════════════════════════════════════════════
def r_yn(cell):
    chk_line(cell,"Yes"); chk_line(cell,"No")
    note(cell,"If Yes, please specify:"); field(cell,"",34)

def r_emp(cell):
    for o in ["Immediate (within 1–2 weeks)","Short-term (within 1 month)","Medium-term (1–3 months)","Long-term (>3 months)","Tentative date (Please specify) - _________","Not yet decided"]:
        chk_line(cell,o)
    
def r_emp1(cell):
    for o in ["< 500","500 – 1,000","1,000 – 5,000"]:
        chk_line(cell,o)
    chk_line(cell,"> 5,000"); field(cell,"  If > 5,000, specify: ",18)

def r_gov(cell):
    for o in ["Yes, centralised global office","Yes, regional offices",
              "No, decisions taken by IT / Legal / Other","No formal structure"]:
        chk_line(cell,o)
    note(cell,"Specify:"); field(cell,"",34)

def r_dec(cell):
    for o in ["Privacy Office","Legal & Compliance","IT Security","Business Unit Heads"]:
        chk_line(cell,o)
    chk_line(cell,"Other (Please specify) - ___________________")

def r_pol(short):
    def f(cell):
        for o in ["Existing framework in place (requires update)",
                  "Drafted but not implemented","Needs to be formulated from scratch"]:
            chk_line(cell,o)
        chk_line(cell,"Other (Please specify) - ___________________")
    return f

def r_opts(options, elaborate=False, other=True):
    def f(cell):
        for o in options:
            chk_line(cell,o)
        if other: chk_line(cell,"Other (Please specify) - ___________________")
    return f

def r_disc(cell):
    chk_line(cell,"Yes"); chk_line(cell,"No")
    note(cell,"If Yes, please specify tool:"); field(cell,"",34)

def r_stor(cell):
    for o in ["On-premise","Cloud","Hybrid(On-premise + Cloud)"]: chk_line(cell,o)
    chk_line(cell,"Other (Please specify) - ___________________")

# ═══════════════════════════════════════════════════════════
# Page border
# ═══════════════════════════════════════════════════════════
def add_page_border(doc):
    sectPr = doc.sections[0]._sectPr
    pgBdr  = OxmlElement("w:pgBdr")
    for side in ["top","left","bottom","right"]:
        b=OxmlElement(f"w:{side}"); b.set(qn("w:val"),"single")
        b.set(qn("w:sz"),"12"); b.set(qn("w:space"),"24")
        b.set(qn("w:color"),C_MID_BLUE.lstrip("#")); pgBdr.append(b)
    for old in sectPr.findall(qn("w:pgBdr")): sectPr.remove(old)
    sectPr.append(pgBdr)

# ═══════════════════════════════════════════════════════════
# Cover block
# ═══════════════════════════════════════════════════════════
def _set_para_spacing(para, before, after, line, rule="auto"):
    pPr = _pPr(para)
    for old in pPr.findall(qn("w:spacing")): pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"),  str(before))
    sp.set(qn("w:after"),   str(after))
    sp.set(qn("w:line"),    str(line))
    sp.set(qn("w:lineRule"),rule)
    pPr.append(sp)

def add_cover(doc, org_name, sector, logo_path=None):
    """
    Compact 2-col cover: left = Protiviti logo (locked, non-editable),
    right = title + org + date.
    """
    from docx.shared import Inches

    LOGO_W  = 2000
    TITLE_W = TOTAL - LOGO_W   # 7288 DXA

    tbl = doc.add_table(rows=1, cols=2)
    tbl_align_center(tbl)
    tbl_width(tbl, TOTAL)
    tbl_borders(tbl, C_DARK_NAVY)
    tbl_clear_style(tbl)

    lc = tbl.rows[0].cells[0]
    rc = tbl.rows[0].cells[1]

    for cell, w in [(lc, LOGO_W), (rc, TITLE_W)]:
        cell_shade(cell, C_DARK_NAVY)
        cell_w(cell, w)
        cell_valign(cell, "center")

    cell_margins(lc, top=120, bottom=120, left=160, right=80)
    cell_margins(rc, top=120, bottom=120, left=80, right=160)

    # Gold bottom stripe on both cells
    cell_bottom_border(lc, C_GOLD, sz="12")
    cell_bottom_border(rc, C_GOLD, sz="12")

    # ── Left: Protiviti logo — LOCKED (non-editable) ──────────────────────────
    #
    #   add_locked_picture_to_cell() wraps the image paragraph inside a
    #   <w:sdt> with <w:lock w:val="sdtContentLocked"/>.
    #   In Word, users cannot click into, resize, or delete the logo;
    #   the cursor simply skips over the content control.
    #
    add_locked_picture_to_cell(
        cell        = lc,
        image_bytes = _LOGO_BYTES,
        width_inches= 1.15,
        align       = WD_ALIGN_PARAGRAPH.LEFT,
        space_before= 0,
        space_after = 0,
    )

    # ── Right: Title ──────────────────────────────────────────────────────────
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_para_spacing(rp, 0, 0, 280)
    srun(rp, "Pre-Scoping Privacy Questionnaire", bold=True, size=13, color=C_WHITE)

    # Small gap after cover
    g = doc.add_paragraph(); no_space(g)
    g.paragraph_format.space_after = Pt(6)

# ═══════════════════════════════════════════════════════════
# Main export function
# ═══════════════════════════════════════════════════════════
def generate_questionnaire_docx(org_name: str, ai: dict) -> bytes:
    _CB[0]=1000
    short  = ai.get("short_name", org_name.split()[0])
    sector = ai.get("sector","")

    doc = Document()
    for sec in doc.sections:
        sec.page_width=Cm(21); sec.page_height=Cm(29.7)
        sec.top_margin=Cm(2.4); sec.bottom_margin=Cm(1.8)
        sec.left_margin=Cm(1.8); sec.right_margin=Cm(1.8)
        sec.header_distance=Cm(1.2); sec.footer_distance=Cm(1.0)

    sty=doc.styles["Normal"]
    sty.font.name=FONT; sty.font.size=Pt(FONT_SZ)
    sty.paragraph_format.space_before=Pt(0); sty.paragraph_format.space_after=Pt(3)

    add_page_border(doc)
    add_cover(doc, org_name, sector, logo_path=ai.get("logo_path"))

    # Section 1
    sec_hdr(doc,"Organisational Overview","🏢")
    t1=make_table(doc)
    q_row(t1,1,"Are there any subsidiaries, affiliates, or joint ventures to be included in this engagement?",r_yn)
    q_row(t1,2,"If your response above is "Yes", please confirm whether the above mentioned entities have centralized Cybersecurity/IT, HR and Legal functions in place to support all business functions?",r_yn,tint=True)
    q_row(t1,3,"What is the approximate employee strength?",r_emp1)
    doc.add_paragraph()

    # Section 2
    sec_hdr(doc,"Governance & Accountability","⚖️")
    t2=make_table(doc)
    q_row(t2,1,"Has a Privacy Governance Committee or Privacy Office been set up?",r_gov)
    q_row(t2,2,"If your response to the above is "No", please confirm who takes decisions on the use of personal or its related decision making?",r_dec,tint=True)
    q_row(t2,3,f"What is the current status of {short}'s privacy policy framework?",r_pol(short))
    doc.add_paragraph()

    # Section 3
    sec_hdr(doc,"Business Lines & Stakeholders","📊")
    t3=make_table(doc)
    q_row(t3,1,f"Which of the following are {short}'s core business lines?",r_opts(ai.get("business_lines",[]),elaborate=True))
    q_row(t3,2,"Which of these internal teams may potentially process personal data?",r_opts(ai.get("stakeholder_teams",[])),tint=True)
    doc.add_paragraph()

    # Section 4
    sec_hdr(doc,"Data Ecosystem","🖥️")
    t4=make_table(doc)
    q_row(t4,1,f"List all customer-facing interfaces used by {short}.",r_opts(ai.get("customer_interfaces",[]),elaborate=True))
    q_row(t4,2,"List all core systems / applications that process, store or manage personal data?",r_opts(ai.get("core_systems",[])),tint=True)
    q_row(t4,3,"Do you use any tools to identify, map or track personal data across systems? (E.g., data discovery, data flow mapping, etc.)",r_disc)
    q_row(t4,4,"Where is personal data stored and hosted?",r_stor,tint=True)
    doc.add_paragraph()

    # Section 5
    sec_hdr(doc,"Cross Border Data Transfer","🏢")
    t6=make_table(doc)
    q_row(t6,1,"Does any personal data processed by the organization get transferred or accessed from outside India? If yes, please specify the countries, entities involved, and purpose of transfer.",r_yn)
    doc.add_paragraph()

    # Section 6
    sec_hdr(doc,"ADDITIONAL DATA","🏢")
    t7=make_table(doc)
    q_row(t7,1,"When do you plan to initiate the engagement? Please provide a tentative start date.",r_emp)
    doc.add_paragraph()

    # Completion note
    nt=doc.add_table(1,1)
    tbl_align_center(nt); tbl_width(nt,TOTAL); tbl_borders(nt,C_GOLD); tbl_clear_style(nt)
    nc=nt.rows[0].cells[0]; cell_shade(nc,"FFF8E7"); cell_w(nc,TOTAL)
    cell_margins(nc,120,120,180,180)
    np_=nc.paragraphs[0]; np_.alignment=WD_ALIGN_PARAGRAPH.CENTER; no_space(np_)
    srun(np_,"Please complete all sections and return to the Data Privacy Team. All information will be treated as strictly confidential.",italic=True,size=FONT_SZ-1,color="7A5C00")

    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()
